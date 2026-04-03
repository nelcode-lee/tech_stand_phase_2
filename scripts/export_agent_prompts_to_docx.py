"""
Export LLM system prompts from src/pipeline/agents/*.py into Word (.docx) files.

Usage (from repository root):
  python scripts/export_agent_prompts_to_docx.py

Output: docs/agent_prompt_exports/*.docx
"""
from __future__ import annotations

import importlib
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

AGENT_MODULES = [
    "cleansing_agent",
    "draft_layout_agent",
    "conflict_agent",
    "specifying_agent",
    "sequencing_agent",
    "terminology_agent",
    "formatting_agent",
    "risk_agent",
    "validation_agent",
]


def _add_paragraphs(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph(line)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)


def _prompt_sections_for_module(mod_name: str, m) -> list[tuple[str, str]]:
    """Return list of (section_title, full_text)."""
    if mod_name == "risk_agent":
        from src.pipeline.agent_rules import (
            CORRECTIVE_ACTIONS_RULE,
            DOCUMENT_REFERENCE_RULE,
            JOB_TITLE_RULE,
            TOLERANCE_VS_REFERENCE_RULE,
        )

        tpl = getattr(m, "_SYSTEM_PROMPT_TEMPLATE", None)
        if not isinstance(tpl, str):
            return []
        full = (
            tpl
            + DOCUMENT_REFERENCE_RULE
            + JOB_TITLE_RULE
            + TOLERANCE_VS_REFERENCE_RULE
            + CORRECTIVE_ACTIONS_RULE
        )
        return [
            ("_SYSTEM_PROMPT_TEMPLATE (base template only)", tpl),
            (
                "Full runtime system string (template + DOCUMENT_REFERENCE_RULE + JOB_TITLE_RULE + TOLERANCE_VS_REFERENCE_RULE + CORRECTIVE_ACTIONS_RULE)",
                full,
            ),
        ]

    sections: list[tuple[str, str]] = []
    names: list[str] = []
    for name in sorted(dir(m)):
        if name.startswith("__"):
            continue
        val = getattr(m, name, None)
        if not isinstance(val, str) or len(val) < 40:
            continue
        if (
            "PROMPT" in name
            or "TEMPLATE" in name
            or name in ("DRAFT_LAYOUT_SYSTEM",)
        ):
            names.append(name)

    for name in names:
        sections.append((name, getattr(m, name)))

    if mod_name == "draft_layout_agent":
        sections.append(
            (
                "Runtime note (not a separate constant)",
                "At runtime, DraftLayoutAgent appends to the system message:\n"
                '  "\\n\\nStandard section names (use where content fits): " + section_names\n'
                "where section_names comes from domain_context.json template plus defaults (see _get_section_names_for_prompt).",
            )
        )

    return sections


def main() -> None:
    out_dir = ROOT / "docs" / "agent_prompt_exports"
    out_dir.mkdir(parents=True, exist_ok=True)

    for mod_name in AGENT_MODULES:
        m = importlib.import_module(f"src.pipeline.agents.{mod_name}")
        sections = _prompt_sections_for_module(mod_name, m)
        if not sections:
            print(f"skip (no prompts found): {mod_name}")
            continue

        doc = Document()
        title = mod_name.replace("_", " ").title()
        h0 = doc.add_heading(f"Agent prompts: {title}", level=0)
        h0.runs[0].font.name = "Calibri"

        intro = doc.add_paragraph(
            "Extracted from the tech_standards_phase2 pipeline. "
            "Strings reflect runtime values where the module concatenates shared rules at import time."
        )
        intro.runs[0].font.size = Pt(10)

        for sec_title, body in sections:
            doc.add_heading(sec_title, level=1)
            _add_paragraphs(doc, body)

        safe = mod_name.replace(".py", "")
        path = out_dir / f"{safe}_prompts.docx"
        doc.save(path)
        print(f"wrote {path}")

    print("done.")


if __name__ == "__main__":
    main()
