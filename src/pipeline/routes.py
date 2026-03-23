"""FastAPI routes for agent pipeline."""
import asyncio
import io
import json
import logging
import re
from collections.abc import Awaitable, Callable
from datetime import datetime, timezone
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import StreamingResponse

log = logging.getLogger(__name__)
from pydantic import BaseModel

from src.pipeline.models import Document, PipelineContext, RequestType, DocLayer
from src.pipeline.router import PipelineRouter
from src.rag.models import DocumentChunk

ProgressEmit = Callable[[dict], Awaitable[None]]


def _agent_to_frontend_step_key(agent_name: str) -> str:
    """Map backend agent / phase name to AnalysePage loading strip `key` values."""
    return {
        "context": "cleansor",
        "cleansing": "cleansor",
        "draft_layout": "specifier",
        "conflict": "conflictor",
        "specifying": "specifier",
        "sequencing": "sequencer",
        "terminology": "terminator",
        "formatting": "formatter",
        "risk": "risk-assessor",
        "validation": "validator",
        "finding_verification": "finding-verifier",
    }.get(agent_name, "cleansor")

router = APIRouter(tags=["pipeline"])


class QueryRequest(BaseModel):
    """Q&A over document library. Returns answer with citations."""
    question: str
    document_id: str | None = None  # Optional: scope to one document
    doc_layer: str | None = None  # Optional: filter by layer (policy, principle, sop, work_instruction)


class GenerateWorkInstructionRequest(BaseModel):
    """Qualifying answers for Work Instruction generation + optional refinement."""
    task_name: str
    parent_sop: str | None = None
    site: str | None = None
    process_type: str | None = None
    has_measurements: bool = False
    measurements_detail: str | None = None
    has_safety: bool = False
    safety_detail: str | None = None
    needs_visuals: bool = False
    needs_checklist: bool = False
    reference_doc_ids: list[str] | None = None
    follow_up_message: str | None = None  # For chat refinement
    previous_draft: str | None = None  # For chat refinement


class AnalyseRequest(BaseModel):
    tracking_id: str
    request_type: str
    doc_layer: str
    sites: list[str] = []
    policy_ref: str | None = None
    document_id: str | None = None  # Optional: from config for metrics
    title: str | None = None  # Optional: document title for metrics
    requester: str | None = None  # Person who requested the analysis (logged with findings)
    attached_doc_url: str | None = None
    content: str | None = None
    retrieved_chunks: list[dict] | None = None
    query: str | None = None  # Optional: semantic search query when using vector retrieval
    agents: list[str] | None = None  # Optional: run only these agents (e.g. for targeted mode)
    additional_doc_ids: list[str] | None = None  # Optional: reference docs to tighten guardrails and find anomalies
    agent_instructions: str | None = None  # Optional: specific knowledge for agents; never supersedes policy


def _to_doc_layer(s: str) -> DocLayer:
    v = (s or "").strip().lower()
    if v in ("policy", "policy_brcgs", "policy_cranswick"):
        return DocLayer.policy
    if v in ("principle", "sop", "work_instruction"):
        return DocLayer(v)
    return DocLayer.sop


def _dedup_key(items: list, key_fn) -> list:
    """Deduplicate by key; keep first occurrence."""
    seen = set()
    out = []
    for x in items:
        k = key_fn(x)
        if k in seen:
            continue
        seen.add(k)
        out.append(x)
    return out


def _deduplicate_findings(ctx: PipelineContext) -> PipelineContext:
    """Remove duplicate findings from chunk overlap or table extraction."""
    def norm(s: str) -> str:
        return (s or "").strip().lower()[:200]

    ctx.risk_gaps = _dedup_key(ctx.risk_gaps, lambda g: norm(g.location) + "|" + norm(g.issue))
    ctx.cleanser_flags = _dedup_key(
        ctx.cleanser_flags,
        lambda f: norm(f.location) + "|" + norm(f.current_text),
    )
    ctx.structure_flags = _dedup_key(ctx.structure_flags, lambda f: norm(f.section) + "|" + norm(f.detail))
    ctx.content_integrity_flags = _dedup_key(
        ctx.content_integrity_flags,
        lambda f: norm(f.location) + "|" + norm(f.excerpt) + "|" + norm(f.detail),
    )
    ctx.specifying_flags = _dedup_key(
        ctx.specifying_flags,
        lambda f: norm(f.location) + "|" + norm(f.current_text),
    )
    ctx.sequencing_flags = _dedup_key(ctx.sequencing_flags, lambda f: norm(f.location) + "|" + norm(f.issue))
    ctx.formatting_flags = _dedup_key(ctx.formatting_flags, lambda f: norm(f.location) + "|" + norm(f.issue))
    ctx.compliance_flags = _dedup_key(ctx.compliance_flags, lambda f: norm(f.location) + "|" + norm(f.issue))
    ctx.terminology_flags = _dedup_key(
        ctx.terminology_flags,
        lambda f: norm(f.term) + "|" + norm(f.location or ""),
    )
    ctx.conflicts = _dedup_key(ctx.conflicts, lambda c: norm(c.description))
    return ctx


def _to_request_type(s: str) -> RequestType:
    valid = ("new_document", "update_existing", "contradiction_flag", "review_request",
             "single_document_review", "harmonisation_review", "principle_layer_review")
    return RequestType(s) if s in valid else RequestType.single_document_review


def _chunks_from_request(req: AnalyseRequest) -> list[DocumentChunk]:
    if req.retrieved_chunks:
        return [
            DocumentChunk(
                text=c.get("text", ""),
                doc_layer=_to_doc_layer(c.get("doc_layer", req.doc_layer)),
                sites=c.get("sites", req.sites),
                policy_ref=c.get("policy_ref") or req.policy_ref,
                document_id=c.get("document_id"),
                source_path=c.get("source_path"),
                title=c.get("title"),
                library=c.get("library"),
                chunk_index=c.get("chunk_index", 0),
            )
            for c in req.retrieved_chunks
        ]
    if req.content:
        return [
            DocumentChunk(
                text=req.content,
                doc_layer=_to_doc_layer(req.doc_layer),
                sites=req.sites,
                policy_ref=req.policy_ref,
                document_id=req.document_id,
                chunk_index=0,
            )
        ]
    # Vector retrieval: fetch relevant chunks from Supabase
    # When document_id is provided, only chunks from that document are returned (scoped analysis)
    from src.rag.retriever import retrieve
    from src.rag.document_registry import get_document_content
    if req.document_id:
        log.info("Retrieving chunks for document_id=%s", req.document_id)
    else:
        log.warning("No document_id in request — retrieval will be unfiltered (all documents)")
    chunks = retrieve(
        doc_layer=req.doc_layer,
        sites=req.sites or None,
        policy_ref=req.policy_ref,
        document_id=req.document_id,
        query_text=req.query,
    )
    # Fallback: when vector store returns no chunks but we have document_id, use full content from registry
    # (handles fresh ingest, indexing delay, or vecs filter mismatch)
    if not chunks and req.document_id:
        try:
            content, _ = get_document_content(req.document_id)
            if content and content.strip():
                log.info("Using document registry content for document_id=%s (vector store had 0 chunks)", req.document_id)
                return [
                    DocumentChunk(
                        text=content,
                        doc_layer=_to_doc_layer(req.doc_layer),
                        sites=req.sites or [],
                        policy_ref=req.policy_ref,
                        document_id=req.document_id,
                        chunk_index=0,
                    )
                ]
        except Exception as e:
            log.warning("Fallback to document registry failed for %s: %s", req.document_id, e)
    return chunks


def _fetch_additional_documents(additional_doc_ids: list[str] | None) -> list[Document]:
    """Fetch content from additional reference documents for guardrails and anomaly detection."""
    if not additional_doc_ids:
        return []
    from src.rag.document_registry import get_document_content
    from src.rag.retriever import retrieve
    from src.rag.models import DocLayer
    sibling_docs: list[Document] = []
    for doc_id in additional_doc_ids:
        doc_id = (doc_id or "").strip()
        if not doc_id:
            continue
        try:
            content, _ = get_document_content(doc_id)
            if content:
                sibling_docs.append(Document(
                    id=doc_id,
                    title=doc_id,
                    content=content[:30000],  # cap per doc
                    doc_layer=DocLayer.sop,
                    sites=[],
                    policy_ref=None,
                ))
            else:
                # Fallback: fetch chunks if full content not stored
                chunks = retrieve(doc_layer=DocLayer.sop, document_id=doc_id, limit=100)
                if chunks:
                    chunks.sort(key=lambda c: c.chunk_index)
                    content = "\n\n".join(c.text for c in chunks)[:30000]
                    sibling_docs.append(Document(id=doc_id, title=doc_id, content=content, doc_layer=DocLayer.sop, sites=[], policy_ref=None))
        except Exception as e:
            log.warning("Could not fetch additional doc %s: %s", doc_id, e)
    return sibling_docs


def _fetch_policy_document(
    *,
    query_text: str,
    document_id: str | None = None,
    standard_name: str | None = None,
) -> Document | None:
    """Fetch one policy document from structured clauses first, then chunk fallback."""
    from src.rag.retriever import retrieve
    from src.rag.models import DocLayer
    from src.rag.document_registry import get_policy_context_block
    try:
        clause_block = ""
        clause_rows: list[dict] = []
        if document_id:
            clause_block, clause_rows = get_policy_context_block(
                document_id=document_id,
                query_text=query_text,
                limit=25,
                max_chars=12000,
            )
        else:
            clause_block, clause_rows = get_policy_context_block(
                standard_name=standard_name,
                query_text=query_text,
                limit=25,
                max_chars=12000,
            )
        if clause_rows and clause_block:
            standard_name = clause_rows[0].get("standard_name") or "Policy"
            version = clause_rows[0].get("version") or ""
            title = f"{standard_name} {version}".strip() + " (relevant clauses)"
            return Document(
                id=document_id or clause_rows[0].get("document_id") or standard_name or "policy",
                title=title,
                content=clause_block,
                doc_layer=DocLayer.policy,
                sites=[],
                policy_ref=document_id,
            )
        if standard_name and not document_id:
            return None

        policy_chunks = retrieve(
            doc_layer=DocLayer.policy,
            document_id=document_id,
            limit=150,
        )
        if not policy_chunks:
            return None
        # Sort by document_id then chunk_index for coherent order
        policy_chunks.sort(key=lambda c: ((c.document_id or ""), c.chunk_index))
        content = "\n\n".join(c.text for c in policy_chunks)
        # Build title from unique document_ids present
        doc_ids = list(dict.fromkeys(c.document_id for c in policy_chunks if c.document_id))
        title = doc_ids[0] if len(doc_ids) == 1 else f"Policy layer ({len(doc_ids)} documents)"
        return Document(
            id=doc_ids[0] if doc_ids else "policy",
            title=title,
            content=content[:50000],  # cap for LLM context
            doc_layer=DocLayer.policy,
            sites=[],
            policy_ref=document_id,
        )
    except Exception as e:
        log.warning("Failed to fetch policy document: %s", e)
        return None


def _fetch_parent_policies(req: AnalyseRequest) -> tuple[Document | None, list[Document]]:
    """
    Fetch layered policies for analysis.

    For SOP / work instruction / principle documents, parent context is always loaded from
    both ingested standards (Cranswick Manufacturing Standard + BRCGS Food Safety). No manual
    policy picker — request policy_ref and library policy_ref on the document are ignored for that path.

    Non-SOP layers still use explicit policy_ref or registry policy_ref when set.
    """
    from src.rag.document_registry import get_document_content, get_document_policy_ref

    is_harmonisation = _to_request_type(req.request_type) == RequestType.harmonisation_review
    is_single_review = _to_request_type(req.request_type) == RequestType.single_document_review
    dl = (req.doc_layer or "").strip().lower()
    use_layered_standards = dl in ("sop", "work_instruction", "principle")

    policy_ref = (req.policy_ref or "").strip()
    if use_layered_standards:
        policy_ref = ""
    elif not policy_ref and req.document_id:
        policy_ref = (get_document_policy_ref(req.document_id) or "").strip()
    has_policy_ref = bool(policy_ref)
    needs_policy = is_harmonisation or has_policy_ref
    if not needs_policy and dl in ("sop", "work_instruction", "principle"):
        needs_policy = True
    if not needs_policy and is_single_review and req.document_id:
        needs_policy = True
    if not needs_policy:
        return None, []

    query_text = (req.query or "").strip()
    if not query_text and req.document_id:
        try:
            current_content, _ = get_document_content(req.document_id)
            query_text = (current_content or "")[:5000]
        except Exception:
            query_text = ""

    parent_policy: Document | None = None
    higher_order_policies: list[Document] = []

    if dl in ("sop", "work_instruction", "principle"):
        # Always both standards (relevant clauses), when present in policy_clause_records / vectors
        parent_policy = _fetch_policy_document(
            query_text=query_text,
            standard_name="Cranswick Manufacturing Standard",
        )
        higher = _fetch_policy_document(
            query_text=query_text,
            standard_name="BRCGS Food Safety",
        )
        if higher and (not parent_policy or higher.id != parent_policy.id):
            higher_order_policies.append(higher)
        return parent_policy, higher_order_policies

    # Non-SOP fallback: explicit policy ref first, otherwise BRCGS only.
    if has_policy_ref:
        parent_policy = _fetch_policy_document(query_text=query_text, document_id=policy_ref)
    if not parent_policy:
        parent_policy = _fetch_policy_document(query_text=query_text, standard_name="BRCGS Food Safety")
    return parent_policy, higher_order_policies


# Map API result keys -> agent names for agent_findings
_FINDING_KEYS_TO_AGENT = {
    "risk_gaps": "risk",
    "cleanser_flags": "cleansing",
    "content_integrity_flags": "cleansing",
    "structure_flags": "cleansing",
    "conflicts": "conflict",
    "specifying_flags": "specifying",
    "terminology_flags": "terminology",
    "compliance_flags": "validation",
    "formatting_flags": "formatting",
    "sequencing_flags": "sequencing",
}


def _doc_id_matches(request_id: str, chunk_doc_id: str) -> bool:
    """Strict match: request doc ID matches chunk doc ID or chunk prefix (e.g. FSP003 matches FSP003-VEHICLE-LOADING), never a different doc."""
    req = (request_id or "").strip().upper()
    chunk = (chunk_doc_id or "").strip().upper()
    if not req or not chunk:
        return False
    if req == chunk:
        return True
    # Request "FSP003" matches chunk "FSP003-VEHICLE-LOADING" (same doc, extended id)
    if chunk.startswith(req) and (len(chunk) == len(req) or chunk[len(req)] in "-:_ "):
        return True
    # Request "FSP003 - Vehicle Loading" matches chunk "FSP003" (take base id from request)
    base = req.split()[0] if req else ""
    if base and chunk == base:
        return True
    if base and len(chunk) >= len(base) and chunk.startswith(base):
        return len(chunk) == len(base) or (len(chunk) > len(base) and chunk[len(base)] in "-:_ ")
    return False


def _filter_chunks_by_document(chunks: list[DocumentChunk], document_id: str | None) -> list[DocumentChunk]:
    """When document_id is set, keep only chunks from that document. Reject any mismatched doc IDs."""
    if not document_id or not chunks:
        return chunks
    doc_id = (document_id or "").strip()
    if not doc_id:
        return chunks
    filtered = [c for c in chunks if _doc_id_matches(doc_id, c.document_id or "")]
    # Log contamination if vector store returned wrong-document chunks
    if chunks and not filtered:
        seen_ids = list(dict.fromkeys((c.document_id or "").strip() for c in chunks if (c.document_id or "").strip()))
        log.warning(
            "Document mismatch: requested document_id=%s but vector store returned chunks for %s. "
            "Using registry fallback if available.",
            document_id, seen_ids[:5] or ["unknown"]
        )
    return filtered


async def _execute_analyse(body: AnalyseRequest, progress_emit: ProgressEmit | None = None) -> dict:
    """Build context, run pipeline, return the same JSON dict as the non-streaming /analyse response."""
    try:
        chunks = _chunks_from_request(body)
        chunks = _filter_chunks_by_document(chunks, body.document_id)
        if body.document_id and not chunks:
            raise HTTPException(
                status_code=400,
                detail=f"No content found for document '{body.document_id}'. The document may not be ingested yet, or ingestion may have failed. Try re-uploading the document.",
            )
        doc_id = (body.document_id or "").strip() or (chunks[0].document_id if chunks else "") or ""
        doc_title = (body.title or "").strip() or (chunks[0].title if chunks else "") or doc_id or ""
        agents_override = body.agents if body.agents else None
        parent_policy, higher_order_policies = _fetch_parent_policies(body)
        sibling_docs = _fetch_additional_documents(body.additional_doc_ids)

        # When document_id set: use full document content for excerpts and to avoid chunk overlap duplicates
        # If registry has no content, reconstruct from retrieved chunks so agents can still quote excerpts
        full_content = None
        if doc_id:
            try:
                from src.rag.document_registry import get_document_content
                full_content, _ = get_document_content(doc_id)
            except Exception:
                pass
            if not full_content and chunks:
                chunks_sorted = sorted(chunks, key=lambda c: getattr(c, "chunk_index", 0))
                full_content = "\n\n".join((c.text or "").strip() for c in chunks_sorted if (c.text or "").strip())

        # Prior user feedback for this document (from finding_notes) — checked before reasoning
        prior_feedback: list[dict] = []
        if doc_id:
            try:
                from src.rag.finding_notes import get_relevant_finding_notes
                prior_feedback = get_relevant_finding_notes(doc_id, limit=20)
            except Exception as e:
                log.debug("Could not load prior feedback for pipeline: %s", e)

        # Glossary for all docs (from domain_context.json) — agents use for terminology
        glossary_block: str | None = None
        try:
            from src.pipeline.domain import get_glossary_block, load_domain_context
            glossary_block = get_glossary_block(load_domain_context()) or None
        except Exception as e:
            log.debug("Could not load glossary for pipeline: %s", e)

        ctx = PipelineContext(
            tracking_id=body.tracking_id,
            request_type=_to_request_type(body.request_type),
            doc_layer=_to_doc_layer(body.doc_layer),
            sites=body.sites,
            policy_ref=body.policy_ref,
            attached_doc_url=body.attached_doc_url,
            document_id=doc_id or None,
            document_title=doc_title or None,
            retrieved_chunks=chunks,
            full_document_content=full_content,
            parent_policy=parent_policy,
            higher_order_policies=higher_order_policies,
            sibling_docs=sibling_docs,
            agent_instructions=(body.agent_instructions or "").strip() or None,
            prior_feedback=prior_feedback,
            glossary_block=glossary_block,
        )
        router_instance = PipelineRouter(agents_override=agents_override)

        progress_callback = None
        if progress_emit:
            n_agents = len(router_instance._select_agents(ctx))
            await progress_emit({"type": "start", "total": 1 + n_agents})
            await progress_emit(
                {
                    "type": "progress",
                    "agent": "context",
                    "step_key": _agent_to_frontend_step_key("context"),
                }
            )

            async def progress_callback(agent_name: str) -> None:
                await progress_emit(
                    {
                        "type": "progress",
                        "agent": agent_name,
                        "step_key": _agent_to_frontend_step_key(agent_name),
                    }
                )

        ctx = await router_instance.run(ctx, progress_callback=progress_callback)

        # Deduplicate findings — chunk overlap or table extraction can produce same finding twice
        ctx = _deduplicate_findings(ctx)

        # Cross-check findings vs full document — drop false positives when limits/refs exist nearby (verbatim-checked)
        try:
            from src.pipeline.finding_verification import run_finding_verification

            if progress_emit:
                await progress_emit(
                    {
                        "type": "progress",
                        "agent": "finding_verification",
                        "step_key": _agent_to_frontend_step_key("finding_verification"),
                    }
                )
            await run_finding_verification(ctx)
        except Exception as e:
            log.warning("Finding verification skipped: %s", e)

        # Compliance flags → grounded policy clause links (candidate retrieval + constrained LLM + verify)
        if ctx.compliance_flags:
            from src.pipeline.clause_mapping import (
                enrich_compliance_flags_clause_mapping,
                ensure_compliance_flags_have_clause_mapping,
            )

            try:
                await enrich_compliance_flags_clause_mapping(ctx)
            except Exception as e:
                log.warning("Clause mapping enrichment skipped: %s", e)
            try:
                ensure_compliance_flags_have_clause_mapping(ctx)
            except Exception:
                pass

        # Persist session for dashboard metrics
        from src.rag.analysis_sessions import record_session

        doc_id = body.document_id or (chunks[0].document_id if chunks else "") or ""
        title = body.title or (chunks[0].title if chunks else "") or doc_id or "Unnamed"
        sites_str = ",".join(body.sites) if body.sites else ""
        total_findings = (
            len(ctx.risk_gaps)
            + len(ctx.cleanser_flags)
            + len(ctx.specifying_flags)
            + len(ctx.structure_flags)
            + len(ctx.content_integrity_flags)
            + len(ctx.sequencing_flags)
            + len(ctx.formatting_flags)
            + len(ctx.compliance_flags)
            + len(ctx.terminology_flags)
            + len(ctx.conflicts)
        )
        agent_findings = {}
        for key, agent in _FINDING_KEYS_TO_AGENT.items():
            val = getattr(ctx, key, None)
            count = len(val) if val else 0
            if count > 0:
                agent_findings[agent] = agent_findings.get(agent, 0) + count
        # Glossary candidates: vague terms to add to glossary (route to HITL)
        glossary_candidates = [
            {"term": t.term, "recommendation": t.recommendation}
            for t in ctx.terminology_flags
            if getattr(t, "glossary_candidate", False)
        ]

        analysis_date = datetime.now(timezone.utc).isoformat()
        requester = (body.requester or "").strip()

        response = {
            "tracking_id": ctx.tracking_id,
            "document_id": doc_id,
            "title": title,
            "doc_layer": ctx.doc_layer.value,
            "requester": requester,
            "analysis_date": analysis_date,
            "draft_ready": ctx.draft_ready,
            "draft_content": ctx.draft_content,
            "overall_risk": ctx.overall_risk.value if ctx.overall_risk else None,
            "conflict_count": ctx.conflict_count,
            "blocker_count": ctx.blocker_count,
            "conflicts": [c.model_dump() for c in ctx.conflicts],
            "terminology_flags": [t.model_dump() for t in ctx.terminology_flags],
            "glossary_candidates": glossary_candidates,
            "risk_scores": [r.model_dump() for r in ctx.risk_scores],
            "risk_gaps": [g.model_dump() for g in ctx.risk_gaps],
            "cleanser_flags": [c.model_dump() for c in ctx.cleanser_flags],
            "specifying_flags": [s.model_dump() for s in ctx.specifying_flags],
            "structure_flags": [s.model_dump() for s in ctx.structure_flags],
            "content_integrity_flags": [c.model_dump() for c in ctx.content_integrity_flags],
            "sequencing_flags": [s.model_dump() for s in ctx.sequencing_flags],
            "formatting_flags": [f.model_dump() for f in ctx.formatting_flags],
            "compliance_flags": [c.model_dump() for c in ctx.compliance_flags],
            "warnings": ctx.warnings,
            "errors": [e.model_dump() for e in ctx.errors],
            "agents_run": ctx.agents_run,
        }
        session_saved = False
        try:
            record_session(
                tracking_id=ctx.tracking_id,
                document_id=doc_id,
                title=title,
                requester=requester,
                doc_layer=ctx.doc_layer.value,
                sites=sites_str,
                overall_risk=ctx.overall_risk.value if ctx.overall_risk else None,
                total_findings=total_findings,
                agents_run=ctx.agents_run,
                agent_findings=agent_findings,
                workflow_type="review",
                result_json=response,
            )
            session_saved = True
        except Exception as e:
            log.warning("Failed to persist analysis session for dashboard: %s", e)

        response["session_saved"] = session_saved
        return response
    except HTTPException:
        raise
    except Exception as e:
        log.exception("Analysis failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e)) from e


@router.post("/analyse")
async def post_analyse(body: AnalyseRequest, stream: bool = Query(False, description="Stream NDJSON progress then final result")):
    """Run the agent pipeline. Use stream=true for newline-delimited progress + final payload."""
    if not stream:
        return await _execute_analyse(body, progress_emit=None)

    async def ndjson_generator():
        queue: asyncio.Queue = asyncio.Queue()

        async def emit(msg: dict) -> None:
            await queue.put(msg)

        async def worker() -> None:
            try:
                result = await _execute_analyse(body, progress_emit=emit)
                await queue.put({"type": "complete", "result": result})
            except HTTPException as he:
                await queue.put({"type": "http_error", "status": he.status_code, "detail": he.detail})
            except Exception as e:
                log.exception("Analysis stream failed: %s", e)
                await queue.put({"type": "error", "message": str(e)})
            finally:
                await queue.put(None)

        task = asyncio.create_task(worker())
        try:
            while True:
                msg = await queue.get()
                if msg is None:
                    break
                yield json.dumps(msg, default=str) + "\n"
        finally:
            await task

    return StreamingResponse(ndjson_generator(), media_type="application/x-ndjson")


class DraftRequest(BaseModel):
    content: str
    filename: str = "draft"


class ValidateSolutionRequest(BaseModel):
    """Re-validate a proposed solution against the original excerpt."""
    excerpt: str = ""
    proposed_solution: str = ""


def _text_to_docx(content: str) -> bytes:
    """Convert plain text (markdown-style headings, lists, simple tables) to DOCX bytes."""
    from docx import Document

    doc = Document()
    lines = content.split("\n")

    # Line classification
    BLANK = "blank"
    H1 = "h1"
    H2 = "h2"
    H3 = "h3"
    CAPS_HEADING = "caps"
    BULLET = "bullet"
    NUMBERED = "numbered"
    TABLE_ROW = "table"
    NORMAL = "normal"

    def classify(line: str) -> tuple[str, str]:
        """Return (kind, payload). payload is text without prefix for lists; row cells for table."""
        stripped = line.strip()
        if not stripped:
            return (BLANK, "")
        if re.match(r"^#\s+.+$", stripped):
            return (H1, re.match(r"^#\s+(.+)$", stripped).group(1).strip())
        if re.match(r"^##\s+.+$", stripped):
            return (H2, re.match(r"^##\s+(.+)$", stripped).group(1).strip())
        if re.match(r"^###\s+.+$", stripped):
            return (H3, re.match(r"^###\s+(.+)$", stripped).group(1).strip())
        # ALL CAPS short line as optional heading (e.g. "SCOPE", "REFERENCES")
        if len(stripped) < 120 and stripped.replace(" ", "").replace("&", "").isupper() and any(c.isalpha() for c in stripped):
            return (CAPS_HEADING, stripped)
        # Bullet: - , * , •
        bullet = re.match(r"^[\s]*[-*•]\s+(.*)$", stripped)
        if bullet:
            return (BULLET, bullet.group(1).strip())
        # Numbered: 1. 2. 2a. 2b. or 1) 2)
        numbered = re.match(r"^[\s]*\d+[a-zA-Z]?[.)]\s+(.*)$", stripped)
        if numbered:
            return (NUMBERED, numbered.group(1).strip())
        # Table: contains |...| or tab-separated
        if "|" in stripped and stripped.count("|") >= 2:
            cells = [c.strip() for c in stripped.split("|") if c.strip()]
            if cells:
                return (TABLE_ROW, "|".join(cells))
        if "\t" in stripped:
            cells = [c.strip() for c in stripped.split("\t")]
            if any(c for c in cells):
                return (TABLE_ROW, "\t".join(cells))
        return (NORMAL, stripped)

    i = 0
    while i < len(lines):
        kind, payload = classify(lines[i])
        if kind == BLANK:
            doc.add_paragraph()
            i += 1
            continue
        if kind == H1:
            doc.add_heading(payload, level=1)
            i += 1
            continue
        if kind == H2:
            doc.add_heading(payload, level=2)
            i += 1
            continue
        if kind == H3:
            doc.add_heading(payload, level=3)
            i += 1
            continue
        if kind == CAPS_HEADING:
            doc.add_heading(payload, level=2)
            i += 1
            continue
        if kind == BULLET:
            while i < len(lines):
                k, p = classify(lines[i])
                if k != BULLET:
                    break
                doc.add_paragraph(p, style="List Bullet")
                i += 1
            continue
        if kind == NUMBERED:
            while i < len(lines):
                k, p = classify(lines[i])
                if k != NUMBERED:
                    break
                doc.add_paragraph(p, style="List Number")
                i += 1
            continue
        if kind == TABLE_ROW:
            sep = "|" if "|" in payload else "\t"
            rows = [payload]
            j = i + 1
            while j < len(lines):
                k, p = classify(lines[j])
                if k != TABLE_ROW:
                    break
                rows.append(p)
                j += 1
            cells_per_row = [len(r.split(sep)) for r in rows]
            ncols = max(cells_per_row) if cells_per_row else 1
            nrows = len(rows)
            table = doc.add_table(rows=nrows, cols=ncols)
            table.style = "Table Grid"
            for ri, row_text in enumerate(rows):
                cell_vals = row_text.split(sep)
                for ci, val in enumerate(cell_vals):
                    if ci < ncols:
                        table.rows[ri].cells[ci].text = val.strip()
            doc.add_paragraph()
            i = j
            continue
        # NORMAL
        doc.add_paragraph(payload)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


@router.post("/draft")
async def post_draft(body: DraftRequest):
    """Generate a DOCX file from draft content. Returns the file for download."""
    try:
        docx_bytes = _text_to_docx(body.content or "")
    except Exception as e:
        log.warning("DOCX generation failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))
    filename = (body.filename or "draft").replace(".docx", "") + ".docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/analyse/validate-solution")
async def post_validate_solution(body: ValidateSolutionRequest):
    """Lightweight check: is the proposed solution appropriate for the excerpt? Returns one-sentence feedback."""
    from src.pipeline.llm import completion
    excerpt = (body.excerpt or "").strip()[:500]
    proposed = (body.proposed_solution or "").strip()[:1000]
    if not excerpt or not proposed:
        return {"feedback": "Provide both excerpt and proposed solution to validate."}
    system = (
        "You are a technical standards reviewer. Given an original document excerpt and a proposed replacement/solution, "
        "you MUST reply in one of two ways:\n"
        "(1) If the solution resolves the finding: clearly agree, e.g. 'This resolves the issue.', 'Agreed — this addresses the gap.', or 'Solution is appropriate.'\n"
        "(2) If the solution needs improvement: give specific further suggestions, e.g. 'Suggest improvement: [specific correction with proper spelling and spacing].' "
        "When suggesting corrections, use proper spelling and spaces between words so the user can copy the text. Be concise and factual."
    )
    prompt = f"Original excerpt:\n{excerpt}\n\nProposed solution:\n{proposed}\n\nDoes this solution resolve the finding? Reply with either agreement or 'Suggest improvement: ...' with specific correction:"
    try:
        feedback = await completion(prompt, system=system)
        return {"feedback": (feedback or "").strip() or "No feedback generated."}
    except Exception as e:
        log.warning("validate-solution failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/analysis/sessions")
async def get_analysis_sessions(limit: int = Query(50, ge=1, le=200)):
    """Return recent analysis sessions for dashboard metrics."""
    from src.rag.analysis_sessions import list_sessions
    return list_sessions(limit=limit)


@router.get("/analysis/sessions/{tracking_id}")
async def get_analysis_session(tracking_id: str):
    """Return a single analysis session with full result (findings, flags, etc.)."""
    from src.rag.analysis_sessions import get_session
    session = get_session(tracking_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    return session


class GlossaryTermRequest(BaseModel):
    """Add a term to the standard glossary (from HITL when terminology is vague)."""
    abbreviation: str
    definition: str


@router.post("/glossary/terms")
async def add_glossary_term_route(body: GlossaryTermRequest):
    """Add a term to the standard glossary. Used when HITL reviewer approves a vague term for inclusion."""
    from src.pipeline.domain import add_glossary_term
    ok = add_glossary_term(body.abbreviation, body.definition)
    if ok:
        return {"ok": True, "message": f"Added '{body.abbreviation}' to glossary"}
    return {"ok": False, "message": f"Term '{body.abbreviation}' already exists in glossary"}


class SaveAnalysisRequest(BaseModel):
    """Save or update an analysis session (e.g. after user edits)."""
    tracking_id: str
    document_id: str = ""
    title: str = ""
    requester: str = ""
    doc_layer: str = "sop"
    sites: str = ""
    overall_risk: str | None = None
    total_findings: int = 0
    agents_run: list[str] = []
    agent_findings: dict = {}
    corrections_implemented: int = 0
    result_json: dict = {}


class FindingNoteAttachment(BaseModel):
    name: str
    contentType: str = "application/octet-stream"
    dataBase64: str


class FindingNoteRequest(BaseModel):
    """Add a user note to a finding. Logged and fed into knowledge base."""
    user_name: str = ""
    document_id: str = ""
    tracking_id: str = ""
    finding_id: str
    finding_summary: dict = {}
    agent_key: str = ""
    note: str
    attachments: list[FindingNoteAttachment] | None = None


class InteractionLogRequest(BaseModel):
    """Store a user interaction event for governance and audit trails."""
    user_name: str = ""
    action_type: str
    route: str = ""
    workflow_mode: str = ""
    document_id: str = ""
    tracking_id: str = ""
    finding_id: str = ""
    doc_layer: str = ""
    metadata: dict = {}


@router.get("/analysis/finding-notes")
async def list_finding_notes_route(limit: int = Query(100, ge=1, le=500)):
    """Return recent user finding notes (logs view)."""
    from src.rag.finding_notes import list_finding_notes
    return list_finding_notes(limit=limit)


@router.post("/analysis/finding-notes")
async def add_finding_note_route(body: FindingNoteRequest):
    """Store a user note on a finding. Logs to DB and adds to vector store for retrieval."""
    from src.rag.finding_notes import add_finding_note
    attachments = [{"name": a.name, "contentType": a.contentType, "dataBase64": a.dataBase64} for a in (body.attachments or [])]
    result = add_finding_note(
        user_name=body.user_name,
        document_id=body.document_id,
        tracking_id=body.tracking_id,
        finding_id=body.finding_id,
        finding_summary=body.finding_summary,
        agent_key=body.agent_key,
        note=body.note,
        attachments=attachments,
        add_to_vector_store=True,
    )
    if not result:
        raise HTTPException(status_code=400, detail="Could not save finding note (empty note?)")
    return {"ok": True, "note": result}


@router.get("/analysis/interaction-logs")
async def list_interaction_logs_route(limit: int = Query(200, ge=1, le=1000)):
    """Return recent governance interaction logs."""
    from src.rag.interaction_logs import list_interaction_logs
    return list_interaction_logs(limit=limit)


@router.post("/analysis/interaction-logs")
async def add_interaction_log_route(body: InteractionLogRequest):
    """Store a governance interaction log entry."""
    from src.rag.interaction_logs import add_interaction_log
    result = add_interaction_log(
        user_name=body.user_name,
        action_type=body.action_type,
        route=body.route,
        workflow_mode=body.workflow_mode,
        document_id=body.document_id,
        tracking_id=body.tracking_id,
        finding_id=body.finding_id,
        doc_layer=body.doc_layer,
        metadata=body.metadata or {},
    )
    if not result:
        raise HTTPException(status_code=400, detail="Could not save interaction log")
    return {"ok": True, "log": result}


@router.post("/analysis/save")
async def save_analysis_session(body: SaveAnalysisRequest):
    """Persist analysis session (captures user changes / ensures state is saved)."""
    from src.rag.analysis_sessions import record_session
    try:
        record_session(
            tracking_id=body.tracking_id,
            document_id=body.document_id or "",
            title=body.title or body.document_id or "Unnamed",
            requester=body.requester or "",
            doc_layer=body.doc_layer or "sop",
            sites=body.sites,
            overall_risk=body.overall_risk,
            total_findings=body.total_findings,
            agents_run=body.agents_run,
            agent_findings=body.agent_findings,
            workflow_type="review",
            corrections_implemented=body.corrections_implemented or 0,
            result_json=body.result_json or {},
        )
        return {"ok": True, "message": "Changes saved"}
    except Exception as e:
        return {"ok": False, "message": str(e)}


@router.post("/query")
async def post_query(body: QueryRequest):
    """
    Q&A over the document library. Retrieves relevant chunks, builds context, and returns
    an answer with citations. Optionally scope to document_id or filter by doc_layer.
    """
    from src.rag.retriever import retrieve
    from src.pipeline.llm import completion

    question = (body.question or "").strip()
    if not question:
        raise HTTPException(status_code=400, detail="question is required")

    doc_layer = body.doc_layer if body.doc_layer in ("policy", "principle", "sop", "work_instruction") else None
    chunks = retrieve(
        doc_layer=doc_layer,
        document_id=body.document_id,
        query_text=question,
        limit=15,
    )
    if not chunks:
        return {"answer": "No relevant documents were found for your question. Try rephrasing or broadening your query.", "citations": []}

    context_parts = []
    seen = set()
    for i, c in enumerate(chunks):
        key = (c.document_id or "", c.chunk_index or 0)
        if key in seen:
            continue
        seen.add(key)
        title = getattr(c, "title", None) or c.document_id or "Document"
        context_parts.append(f"[{i + 1}] ({title}):\n{c.text}")
    context = "\n\n".join(context_parts)

    system = (
        "You are a helpful assistant answering questions about technical standards documents. "
        "Use only the provided context. If the context does not contain enough information, say so. "
        "Cite sources by number (e.g. [1], [2]). Be concise and accurate."
    )
    prompt = f"Context:\n\n{context}\n\nQuestion: {question}\n\nAnswer:"

    try:
        answer = await completion(prompt, system=system)
    except Exception as e:
        log.warning("Query LLM failed: %s", e)
        raise HTTPException(status_code=500, detail="Failed to generate answer")

    seen_ids = set()
    unique_citations = []
    for c in chunks:
        doc_id = c.document_id or ""
        if doc_id and doc_id not in seen_ids:
            seen_ids.add(doc_id)
            unique_citations.append({
                "document_id": doc_id,
                "title": getattr(c, "title", None) or doc_id or "Document",
            })

    return {"answer": answer, "citations": unique_citations}


@router.post("/analysis/generate-work-instruction")
async def generate_work_instruction_route(body: GenerateWorkInstructionRequest):
    """
    Generate or refine a Work Instruction from qualifying questions.
    First call: use task_name + questionnaire fields.
    Refinement: use follow_up_message + previous_draft.
    """
    from src.pipeline.generate_work_instruction import generate_work_instruction
    from src.rag.document_registry import (
        get_document_content,
        get_policy_clauses,
        distinct_policy_document_ids_for_standard_names,
    )

    task_name = (body.task_name or "").strip()
    if not task_name and not body.follow_up_message:
        raise HTTPException(status_code=400, detail="task_name is required for initial generation")

    # Fetch policy context (BRCGS + Cranswick MS)
    from src.pipeline.clause_mapping import _PINNED_POLICY_DOCUMENT_IDS
    doc_ids = distinct_policy_document_ids_for_standard_names(
        ["BRCGS Food Safety", "Cranswick Manufacturing Standard"],
        extra_document_ids=_PINNED_POLICY_DOCUMENT_IDS,
    )
    policy_clauses: list[dict] = []
    for did in doc_ids[:2]:
        policy_clauses.extend(get_policy_clauses(document_id=did, limit=30))
    if not policy_clauses and task_name:
        policy_clauses = get_policy_clauses(standard_name="BRCGS Food Safety", limit=20)
        policy_clauses.extend(get_policy_clauses(standard_name="Cranswick Manufacturing Standard", limit=20))

    # Fetch reference docs if provided
    ref_contents: list[str] = []
    for doc_id in (body.reference_doc_ids or []):
        doc_id = str(doc_id).strip()
        if not doc_id:
            continue
        try:
            content, _ = get_document_content(doc_id)
            if content and content.strip():
                ref_contents.append(content)
        except Exception as e:
            log.debug("Could not fetch ref doc %s: %s", doc_id, e)

    draft, suggested_id = await generate_work_instruction(
        task_name=task_name or "Untitled",
        parent_sop=body.parent_sop,
        site=body.site,
        process_type=body.process_type,
        has_measurements=body.has_measurements,
        measurements_detail=body.measurements_detail,
        has_safety=body.has_safety,
        safety_detail=body.safety_detail,
        needs_visuals=body.needs_visuals,
        needs_checklist=body.needs_checklist,
        reference_doc_contents=ref_contents if ref_contents else None,
        follow_up_message=body.follow_up_message,
        previous_draft=body.previous_draft,
        policy_clauses=policy_clauses,
    )
    return {"draft": draft, "suggested_document_id": suggested_id}
