"""
DOCX export for analysis results — used by POST /analysis/audit-pack.docx.

Builds a structured Word document (headings, tables, bullet lists) from the same
analysis result dict shape returned by POST /analyse.
"""
from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Inches

from src.pipeline.audit_report_export import (
    effective_hazard_control_for_risk_gap,
    hazard_control_label,
    stable_finding_id,
)


def _sites_display(data: dict) -> str:
    sites = data.get("sites")
    if isinstance(sites, list) and sites:
        return ", ".join(str(s) for s in sites)
    if isinstance(sites, str) and sites.strip():
        return sites.strip()
    return "—"


def _safe_str(v) -> str:
    return str(v) if v is not None else "—"


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        r = table.add_row().cells
        r[0].text = k
        r[1].text = v
    doc.add_paragraph("")


def _add_count_table(doc: Document, rows: list[tuple[str, int]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Count"
    for k, n in rows:
        r = table.add_row().cells
        r[0].text = k
        r[1].text = str(int(n or 0))
    doc.add_paragraph("")


def _add_bullets(doc: Document, items: list[str]) -> None:
    for s in items:
        p = doc.add_paragraph(s or "—")
        p.style = "List Bullet"


def export_docx_bytes(data: dict, *, audit_pack: bool = True) -> bytes:
    """
    Return DOCX file bytes for the analysis result dict.
    """
    doc = Document()

    title = "Audit & regulatory readiness pack" if audit_pack else "Analysis report"
    doc.add_heading(title, level=0)

    subtitle = (
        "Automated analysis summary for audit preparation. Policy clause links are shown where verified against ingested standard text."
        if audit_pack
        else "Automated analysis summary."
    )
    doc.add_paragraph(subtitle)

    doc_id = _safe_str(data.get("document_id") or "—")
    doc_title = _safe_str(data.get("title") or "—")
    requester = _safe_str(data.get("requester") or "—")
    analysis_date = _safe_str(data.get("analysis_date") or "—")
    if analysis_date and analysis_date != "—":
        try:
            dt = datetime.fromisoformat(str(analysis_date).replace("Z", "+00:00"))
            analysis_date = dt.strftime("%Y-%m-%d %H:%M UTC")
        except Exception:
            pass

    _add_kv_table(
        doc,
        [
            ("Document", f"{doc_id} — {doc_title}"),
            ("Document layer", _safe_str(data.get("doc_layer") or "—")),
            ("Sites", _sites_display(data)),
            ("Requester", requester),
            ("Analysis date", analysis_date),
            ("Tracking ID", _safe_str(data.get("tracking_id") or "—")),
            ("Draft ready", _safe_str(data.get("draft_ready", "—"))),
            ("Overall risk", _safe_str((data.get("overall_risk") or "—")).upper() if data.get("overall_risk") else "—"),
            ("Agents run", ", ".join(data.get("agents_run", []) or []) or "—"),
            ("Conflicts", _safe_str(data.get("conflict_count", len(data.get("conflicts", []) or [])))),
            ("Blockers", _safe_str(data.get("blocker_count", 0))),
        ],
    )

    if audit_pack:
        fd = data.get("finding_dispositions") if isinstance(data.get("finding_dispositions"), dict) else {}
        fgn = data.get("finding_governance_notes") if isinstance(data.get("finding_governance_notes"), dict) else {}
        fht = data.get("finding_hazard_control_tags") if isinstance(data.get("finding_hazard_control_tags"), dict) else {}
        risk_hz_model: dict[str, str] = {}
        for _g in data.get("risk_gaps") or []:
            if isinstance(_g, dict):
                rid = stable_finding_id("risk", _g)
                risk_hz_model[rid] = str(_g.get("hazard_control_type") or "").strip()
        row_ids = sorted(set(fd.keys()) | set(fgn.keys()) | set(fht.keys()), key=lambda x: str(x))
        if row_ids:
            doc.add_heading("Finding disposition (QA)", level=1)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Finding ID"
            hdr[1].text = "Disposition"
            hdr[2].text = "Governance note"
            hdr[3].text = "Hazard control"
            for fid in row_ids:
                disp = fd.get(fid, "")
                note = (fgn.get(fid) or "").strip()
                if isinstance(disp, dict):
                    disp = str(disp)
                hz_raw = str(fht.get(fid) or "").strip() or risk_hz_model.get(str(fid), "")
                hz = hazard_control_label(hz_raw) or "—"
                r = table.add_row().cells
                r[0].text = str(fid)
                r[1].text = disp or "—"
                r[2].text = note or "—"
                r[3].text = hz
            doc.add_paragraph("")

    # Flag summary
    doc.add_heading("Flag Summary", level=1)
    count_rows = [
        ("Conflicts", len(data.get("conflicts", []) or [])),
        ("Terminology flags", len(data.get("terminology_flags", []) or [])),
        ("Risk gaps", len(data.get("risk_gaps", []) or [])),
        ("Cleanser flags", len(data.get("cleanser_flags", []) or [])),
        ("Specifying flags", len(data.get("specifying_flags", []) or [])),
        ("Structure flags", len(data.get("structure_flags", []) or [])),
        ("Content integrity flags", len(data.get("content_integrity_flags", []) or [])),
        ("Sequencing flags", len(data.get("sequencing_flags", []) or [])),
        ("Formatting flags", len(data.get("formatting_flags", []) or [])),
        ("Compliance flags", len(data.get("compliance_flags", []) or [])),
        ("Errors", len(data.get("errors", []) or [])),
        ("Warnings", len(data.get("warnings", []) or [])),
    ]
    _add_count_table(doc, count_rows)

    # Optional: agent timings
    timings = data.get("agent_timings") or []
    if isinstance(timings, list) and timings:
        doc.add_heading("Agent Timings", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Agent"
        hdr[1].text = "Duration (ms)"
        for t in timings:
            if not isinstance(t, dict):
                continue
            r = table.add_row().cells
            r[0].text = _safe_str(t.get("agent") or "—")
            r[1].text = _safe_str(t.get("duration_ms") or 0)
        doc.add_paragraph("")

    # Conflicts
    conflicts = data.get("conflicts", []) or []
    doc.add_heading("Conflicts", level=1)
    doc.add_paragraph(f"Count: {len(conflicts)}")
    for i, c in enumerate(conflicts, 1):
        if not isinstance(c, dict):
            continue
        doc.add_heading(f"Conflict {i}", level=2)
        _add_kv_table(
            doc,
            [
                ("Type", _safe_str(c.get("conflict_type", ""))),
                ("Severity", _safe_str(c.get("severity", "")).upper()),
                ("Layer", _safe_str(c.get("layer", ""))),
                ("Sites", ", ".join(c.get("sites", []) or []) or "—"),
                ("Documents", ", ".join(c.get("document_refs", []) or []) or "—"),
                ("Blocks Draft", _safe_str(c.get("blocks_draft", False))),
            ],
        )
        desc = _safe_str(c.get("description", "")).strip() or "—"
        doc.add_paragraph(desc)
        rec = _safe_str(c.get("recommendation", "")).strip()
        if rec:
            doc.add_paragraph(f"Recommendation: {rec}")

    # Risk gaps (top by score)
    gaps = data.get("risk_gaps", []) or []
    fht_gaps = data.get("finding_hazard_control_tags") if isinstance(data.get("finding_hazard_control_tags"), dict) else {}
    doc.add_heading("Risk Gaps", level=1)
    doc.add_paragraph(f"Count: {len(gaps)}")
    gaps_sorted = sorted(
        [g for g in gaps if isinstance(g, dict)],
        key=lambda g: g.get("fmea_score", 0) or 0,
        reverse=True,
    )
    for g in gaps_sorted[:50]:
        doc.add_heading(_safe_str(g.get("location", "—")), level=2)
        rpn_lines = []
        if g.get("fmea_band"):
            rpn_lines.append(f"HACCP RPN band: {str(g.get('fmea_band')).upper()}")
        if g.get("fmea_score"):
            rpn_lines.append(f"RPN score: {_safe_str(g.get('fmea_score'))}")
        dims = []
        if g.get("severity"):
            dims.append(f"Severity={g.get('severity')}")
        lik = g.get("likelihood") or g.get("scope")
        if lik:
            dims.append(f"Likelihood={lik}")
        if g.get("detectability"):
            dims.append(f"Detectability={g.get('detectability')}")
        if dims:
            rpn_lines.append("Dimensions: " + " × ".join(dims))
        if rpn_lines:
            doc.add_paragraph(" | ".join(rpn_lines))
        hz_l = hazard_control_label(effective_hazard_control_for_risk_gap(g, fht_gaps))
        if hz_l:
            doc.add_paragraph(f"Hazard control (CCP / oPRP / PRP): {hz_l}")
        issue = _safe_str(g.get("issue", "")).strip()
        if issue:
            doc.add_paragraph(f"Issue: {issue}")
        risk = _safe_str(g.get("risk", "")).strip()
        if risk:
            doc.add_paragraph(f"Risk: {risk}")
        rec = _safe_str(g.get("recommendation", "")).strip()
        if rec:
            doc.add_paragraph(f"Recommendation: {rec}")

    # Compliance flags (table)
    compliance = data.get("compliance_flags", []) or []
    doc.add_heading("Compliance Flags", level=1)
    doc.add_paragraph(f"Count: {len(compliance)}")
    if compliance:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Location"
        hdr[1].text = "Issue"
        hdr[2].text = "Recommendation"
        hdr[3].text = "Policy clause mapping"
        for f in compliance[:80]:
            if not isinstance(f, dict):
                continue
            cm = f.get("clause_mapping") if isinstance(f.get("clause_mapping"), dict) else {}
            status = _safe_str((cm.get("status") or "unmapped")).strip()
            cite = _safe_str(cm.get("canonical_citation") or cm.get("clause_id") or "—")
            mapping = f"{status}: {cite}" if cite else status
            r = table.add_row().cells
            r[0].text = _safe_str(f.get("location") or "—")
            r[1].text = _safe_str(f.get("issue") or "—")
            r[2].text = _safe_str(f.get("recommendation") or "—")
            r[3].text = mapping
        doc.add_paragraph("")

    # Appendices: keep minimal in v1 but ensure the output is a structured doc.
    doc.add_paragraph("")
    doc.add_heading("Appendix", level=1)
    doc.add_paragraph("Generated by Technical Standards AI Platform.")

    # Reasonable default margins via section properties (Word defaults are fine); add page width hint if needed.
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

