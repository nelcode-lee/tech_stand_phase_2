"""Extract plain text from DOCX files."""
from io import BytesIO

try:
    from docx import Document
except ImportError:
    Document = None


def _is_list_paragraph(para) -> bool:
    """Return True if paragraph has list/bullet styling (bullets are not in p.text)."""
    try:
        name = (para.style.name or "").lower()
        return any(x in name for x in ("list", "bullet", "bullrt", "number"))
    except Exception:
        return False


def extract_text_from_docx(file_bytes: bytes) -> str | None:
    """
    Extract plain text from a DOCX file. List items get a bullet prefix so extraction preserves structure.
    Returns None if extraction fails or python-docx is not installed.
    """
    if Document is None:
        return None
    try:
        doc = Document(BytesIO(file_bytes))
        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            if _is_list_paragraph(p):
                paragraphs.append("• " + text)
            else:
                paragraphs.append(text)
        for table in doc.tables:
            paragraphs.append("[TABLE]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    paragraphs.append("\t".join(cells))
        return "\n\n".join(paragraphs) if paragraphs else None
    except Exception:
        return None
